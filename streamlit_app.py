import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
import io


def docx_to_bytes(doc_object):
    bytes_io = io.BytesIO()
    doc_object.save(bytes_io)
    bytes_io.seek(0)
    return bytes_io.getvalue()

def remove_elements_recursive(element, ns):
    # Process current element's children first
    for child in list(element):  # Create a list copy since we'll modify during iteration
        # Check if this child is tcW or tblGrid
        if child.tag == f'{{{ns["w"]}}}tcW' or child.tag == f'{{{ns["w"]}}}tblW' or child.tag == f'{{{ns["w"]}}}tblGrid':
            element.remove(child)
        else:
            # Recursively process this child's children
            remove_elements_recursive(child, ns)

def clear_cell_widths(table, depth):
    if depth < 0:
        return

    for row in table.rows:
        for cell in row.cells:
            # Clear cell width
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find('.//w:tcW', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if tcW is not None:
                tcPr.remove(tcW)
            
            # Handle nested tables
            for nested_table in cell.tables:
                clear_cell_widths(nested_table, depth - 1)

def distribute_columns_evenly(table):
    # Get table width
    tbl = table._tbl
    tblPr = tbl.find('.//w:tblPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # tblW = tblPr.find('.//w:tblW', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    # if tblW is not None:
    #     table_width = int(tblW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0'))
    # else:
    #     # If no width specified, use default page width (about 6 inches in twips)
    #     table_width = 9026  # ~6 inches in twips
    table_width = 11144  # ~7.6 inch

    # Calculate equal width for each column
    num_cols = len(table.columns)
    col_width = int(table_width / num_cols)

    # Set width for each cell in each column
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)

            tcW = tcPr.find('.//w:tcW', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            
            tcW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', str(col_width))
            tcW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')

            # Handle nested tables recursively
            for nested_table in cell.tables:
                distribute_columns_evenly(nested_table)

# Show title and description.
st.title("📄 PHF Editor for JC")
st.write(
    "Upload a document below and ask a question about it – GPT will answer! "
    "To use this app, you need to provide an OpenAI API key, which you can get [here](https://platform.openai.com/account/api-keys). "
)

phf_file = st.file_uploader(
    "Upload PHF document (.doc, .docx)", type=("doc", "docx")
)

if phf_file:
    phf_doc = Document(phf_file)

    # Iterate over physician's orders
    for _, table in enumerate(phf_doc.tables):
        clear_cell_widths(table, 2)
        distribute_columns_evenly(table.cell(0, 0).tables[0].cell(0, 0).tables[0].cell(2, 0).tables[0])

    # Fix lines on the bottom of page

    # Replace "Instructions" to "Recommendations"

    # Extract data to excel sheet with patient info and medications list

    # Let Jon add notes to Recommendations under each patient

    # Save as password protected PDF
    st.download_button(
        label = "Download",
        data = docx_to_bytes(phf_doc),
        file_name = "test.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
